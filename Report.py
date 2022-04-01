import io
import platform
import subprocess

import docx.document
import docx.table
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml import ns
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import _Cell
from docx2pdf import convert
from RepricingStatistics import RepricingStatistics
from RepricingStatisticsOrdinary import RepricingStatisticsOrdinary


class BasicReport:
    def __init__(self):
        pass

    @staticmethod
    def column_text_change(data: pd.DataFrame,
                           table: docx.table.Table,
                           size: int,
                           bold: bool = False):
        for i in range(len(data.columns)):
            run = table.cell(0, i).paragraphs[0].runs[0]
            run.font.size = Pt(size)
            run.font.bold = bold

    @staticmethod
    def cell_text_change(data: pd.DataFrame,
                         table: docx.table.Table,
                         size: int,
                         bold: bool = False):
        for i in range(data.shape[0]):
            for j in range(data.shape[1]):
                run = table.cell(i + 1, j).paragraphs[0].runs[0]
                run.font.size = Pt(size)
                run.font.bold = bold

    @staticmethod
    def delete_columns(table, columns):
        # sort columns descending
        columns.sort(reverse=True)

        grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
        for ci in columns:
            for cell in table.column_cells(ci):
                cell._tc.getparent().remove(cell._tc)

            # Delete column reference.
            col_elem = grid[ci]
            grid.remove(col_elem)

    @staticmethod
    def generate_pdf_windows(doc_path: str, out_path: str):
        """Generate pdf file for windows system"""
        convert(doc_path, out_path)

    @staticmethod
    def generate_pdf_Linux(doc_path, out_path):
        """Generate pdf file for windows system"""
        subprocess.call(['soffice',
                         # '--headless',
                         '--convert-to',
                         'pdf',
                         '--outdir',
                         out_path,
                         doc_path])
        return doc_path

    @staticmethod
    def set_row_height(row, height):
        trPr = row.tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height))
        trPr.append(trHeight)

    @staticmethod
    def set_vertical_cell_direction(cell: _Cell, direction: str):
        # direction: tbRl -- top to bottom, btLr -- bottom to top
        assert direction in ("tbRl", "btLr")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        textDirection = OxmlElement('w:textDirection')
        textDirection.set(qn('w:val'), direction)  # btLr tbRl
        tcPr.append(textDirection)

    @staticmethod
    def create_element(name):
        return OxmlElement(name)

    @staticmethod
    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def create_document(self, header):
        document = Document()
        section = document.sections[-1]
        section.left_martin = Inches(0.1)
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0

        document.add_paragraph().add_run(header).bold = True
        document.add_paragraph(" ")

        section.footer.paragraphs[0].text = header
        section.footer.add_paragraph()
        self.add_page_number(section.footer.paragraphs[1].add_run())
        section.footer.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return document

    def add_page_number(self, run):
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


class OrdinaryReport(BasicReport):
    def __init__(self, header: str, df_cp: pd.DataFrame, target_product: str,
                 main_company: str, period: int, start_date: str, end_date: str, test_start: str, test_end: str):
        super().__init__()
        self.header = header
        self.start_date = start_date
        self.end_date = end_date
        self.interval_information = RepricingStatisticsOrdinary(df_cp,
                                                                target_product,
                                                                main_company,
                                                                period,
                                                                test_start,
                                                                test_end)

    def add_profit_info(self, document: docx.document.Document, data: pd.DataFrame):
        data = data.copy()
        t = document.add_table(data.shape[0] + 1, data.shape[1] + 1)
        t.style = 'TableGrid'
        t.autofit = False
        t.allow_autofit = False
        # add the header rows.
        t.cell(0, 0).text = 'Название колонок'
        t.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # add the rest of the data frame
        for i in range(data.shape[0]):
            t.cell(i + 1, 0).text = str(data.index[i])
            for j, column in enumerate(data.columns):
                if column in ['mean_m_orders', 'median_m_orders', 'mean_price', 'median_price', 'mean_proportion',
                              'profit', 'profit_mean', 'median_proportion']:
                    t.cell(i + 1, j + 1).text = '{:.2f}'.format(data.values[i, j])
                    t.cell(i + 1, j + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    t.cell(i + 1, j + 1).text = str(data.values[i, j])
                    t.cell(i + 1, j + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        hdr_cells = t.rows[0].cells
        hdr_cells[1].text = 'The start of the interval'
        hdr_cells[2].text = 'The end of the interval'
        hdr_cells[3].text = 'Mean m_orders'
        hdr_cells[4].text = 'Median m_orders'
        hdr_cells[5].text = 'Mean price'
        hdr_cells[6].text = 'Median price'
        hdr_cells[7].text = 'Mean proportion'
        hdr_cells[8].text = 'Median proportion'
        hdr_cells[9].text = 'Profit'
        hdr_cells[10].text = 'Profit_mean'
        hdr_cells[11].text = 'Interval'

        for i in range(len(data.columns) + 1):
            super().set_vertical_cell_direction(hdr_cells[i], 'btLr')
            t.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in t.rows[:1]:
            row.height = Inches(1)
            row.width = Inches(1)
        self.delete_columns(t, [0])

        for i in range(len(data.columns)):
            col = t.columns[i]
            col.width = Inches(0.6)

        self.column_text_change(data, t, 12)
        self.cell_text_change(data, t, 9)

    def create_word_file(self):
        document = self.create_document(self.header)
        summary = self.interval_information.summary_calculation()
        self.add_profit_info(document, summary)
        self.interval_information.generate_plots(allow_plot=False)
        graphs = [self.interval_information.mean_orders,
                  self.interval_information.mean_orders_sum,
                  self.interval_information.mean_prop,
                  self.interval_information.prices_df,
                  self.interval_information.mean_profit
                  ]
        for graph in graphs:
            document.add_paragraph(" ")
            memfile = io.BytesIO()
            graph.savefig(memfile, bbox_inches='tight')
            document.add_picture(memfile, width=Inches(6.4), height=Inches(2.25))
            memfile.close()

        document.save('OrdinaryReport.docx')

    def generate_pdf(self, docx_path: str, out_path):
        if platform.system() == 'Linux':
            self.generate_pdf_Linux(docx_path, out_path)
        elif platform.system() == 'Windows':
            self.generate_pdf_windows(docx_path, out_path)


class StatisticalReport(BasicReport):
    def __init__(self, header: str, df_cp: pd.DataFrame, df_cp_best: pd.DataFrame, target_product: str,
                 main_company: str, period: int, start_date: str, end_date: str, test_start: str, test_end: str):
        super().__init__()
        self.header = header
        self.start_date = start_date
        self.end_date = end_date
        self.interval_information = RepricingStatistics(df_cp,
                                                        df_cp_best,
                                                        target_product,
                                                        main_company,
                                                        period,
                                                        test_start,
                                                        test_end)
        # self.metric = Metrics(df_cp_best, test_start)

    def add_profit_info(self, document: docx.document.Document, data: pd.DataFrame):
        data = data.copy()
        data['from'] = data["from"].dt.strftime("%y.%m.%d")
        data['to'] = data["to"].dt.strftime("%y.%m.%d")
        t = document.add_table(data.shape[0] + 1, data.shape[1] + 1)
        t.style = 'TableGrid'
        t.autofit = False
        t.allow_autofit = False
        # add the rest of the data frame
        for i in range(data.shape[0]):
            t.cell(i + 1, 0).text = str(data.index[i])
            for j, column in enumerate(data.columns):
                if column in ['mean_m_orders', 'median_m_orders', 'mean_price', 'median_price', 'mean_proportion',
                              'profit', 'profit_mean', 'median_proportion']:
                    t.cell(i + 1, j + 1).text = '{:.2f}'.format(data.values[i, j])
                    t.cell(i + 1, j + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    t.cell(i + 1, j + 1).text = str(data.values[i, j])
                    t.cell(i + 1, j + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        hdr_cells = t.rows[0].cells
        hdr_cells[1].text = 'The start of the interval'
        hdr_cells[2].text = 'The end of the interval'
        hdr_cells[3].text = 'Mean m_orders'
        hdr_cells[4].text = 'Median m_orders'
        hdr_cells[5].text = 'Mean price'
        hdr_cells[6].text = 'Median price'
        hdr_cells[7].text = 'Mean proportion'
        hdr_cells[8].text = 'Median proportion'
        hdr_cells[9].text = 'Profit'
        hdr_cells[10].text = 'Profit_mean'
        hdr_cells[11].text = 'Interval'

        for i in range(len(data.columns) + 1):
            super().set_vertical_cell_direction(hdr_cells[i], 'btLr')
            t.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in t.rows[:1]:
            row.height = Inches(1)
            row.width = Inches(1)
        self.delete_columns(t, [0])

    def create_word_file(self):
        document = self.create_document(self.header)
        summary = self.interval_information.summary_calculation()
        self.add_profit_info(document, summary)
        self.interval_information.generate_plots(allow_plot=False)
        graphs = [self.interval_information.mean_orders,
                  self.interval_information.mean_orders_best,
                  self.interval_information.mean_orders_sum,
                  self.interval_information.mean_prop,
                  self.interval_information.mean_prop_best,
                  self.interval_information.prices_df,
                  self.interval_information.mean_profit
                  ]
        for graph in graphs:
            document.add_paragraph(" ")
            memfile = io.BytesIO()
            graph.savefig(memfile, bbox_inches='tight')
            document.add_picture(memfile, width=Inches(6.4), height=Inches(2.25))
            memfile.close()

        # document.add_paragraph(" ")
        # self.metric.fit_gam(max_iter=[50, 500, 50], lam=[1, 1000, 100])
        # document.add_paragraph("Scores_cv: ")
        # document.add_paragraph(f"\tR2 score: {np.round(self.metric.result['m_orders']['scores_cv'][0]['R2'], 3)}")
        # document.add_paragraph(f"\tMAE score: {np.round(self.metric.result['m_orders']['scores_cv'][0]['MAE'], 3)}")
        # document.add_paragraph(f"\tWape score: {np.round(self.metric.result['m_orders']['scores_cv'][0]['Wape'], 3)}")
        # document.add_paragraph(f"\tRMSE score: {np.round(self.metric.result['m_orders']['scores_cv'][0]['RMSE'], 3)}")
        # document.add_paragraph("Scores_insample: ")
        # document.add_paragraph(f"\tR2: {np.round(self.metric.result['m_orders']['scores_insample'][0]['R2'], 3)}")
        # document.add_paragraph(f"\tMAE: {np.round(self.metric.result['m_orders']['scores_insample'][0]['MAE'], 3)}")
        # document.add_paragraph(f"\tWape: {np.round(self.metric.result['m_orders']['scores_insample'][0]['Wape'], 3)}")
        # document.add_paragraph(f"\tRMSE: {np.round(self.metric.result['m_orders']['scores_insample'][0]['RMSE'], 3)}")

        document.save('StatisticalReport.docx')

    def generate_pdf(self, docx_path: str, out_path):
        if platform.system() == 'Linux':
            self.generate_pdf_Linux(docx_path, out_path)
        elif platform.system() == 'Windows':
            self.generate_pdf_windows(docx_path, out_path)
